#!/usr/bin/env python3
"""
Build Attira_Growth_Strategy.docx — the founder-facing growth playbook.

Explains, in plain language: the goal (2,000 waitlist signups before launch),
the North Star metric, the funnel, the weekly KPIs, the per-platform roadmap,
how to read the companion dashboard, and what Phase 2 (beta tracking) adds.

Run:  .venv-growth/bin/python scripts/build_strategy_doc.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BRAND = RGBColor(0x1D, 0x4A, 0xFF)
GREY = RGBColor(0x66, 0x66, 0x66)
DARK = RGBColor(0x22, 0x22, 0x22)

doc = Document()

# ── base styles ────────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for lvl, size in [("Heading 1", 16), ("Heading 2", 13), ("Title", 26)]:
    st = doc.styles[lvl]
    st.font.color.rgb = BRAND
    st.font.size = Pt(size)
    st.font.bold = True


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def para(text, italic=False, bold=False, color=None, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(head)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


# ══ COVER ══════════════════════════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("Attira — Pre-Launch Growth Plan")
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = BRAND
para("Social-media roadmap, North Star metric & weekly KPIs on the road to 2,000 waitlist signups",
     italic=True, color=GREY, size=12)
para("Prepared for the founder  •  Companion file: Attira_Growth_Dashboard.xlsx", color=GREY, size=10)
doc.add_paragraph()

# ══ 1. EXECUTIVE SUMMARY ═══════════════════════════════════════════════════════
h1("1.  Executive summary")
para("We are about to push Attira across Reddit, X (Twitter), Threads, Facebook, LinkedIn, "
     "Instagram and Product Hunt, all driving to our live website. This document defines how "
     "we will know — every single week — whether that push is actually working.")
para("The goal for this phase is one number:", bold=True)
para("Reach at least 2,000 waitlist signups before the official App Store launch.",
     bold=True, color=BRAND, size=14)
para("Because we are not yet handing out the app by link, our North Star metric for now is "
     "Waitlist signups per week. Everything we measure ladders up to whether we are on pace "
     "to hit 2,000 in time. Once we start distributing the beta (TestFlight) — roughly two "
     "weeks out — we add a second layer of tracking (Phase 2, section 7).")

# ══ 2. THE FUNNEL & METRICS ════════════════════════════════════════════════════
h1("2.  How we measure: the funnel")
para("Every signup travels the same path. We measure each step so we can see exactly where "
     "people drop off and fix that step.")
para("Social impressions  →  Engagement  →  Link clicks  →  Website visitors  →  WAITLIST SIGNUP",
     bold=True, color=BRAND)
para("(Phase 2, later:  →  Beta link click  →  TestFlight install  →  Activated user)",
     italic=True, color=GREY)

h2("North Star & the supporting KPIs")
table(
    ["KPI", "What it tells us", "Where the number comes from"],
    [
        ["Waitlist signups / week  (NORTH STAR)", "Real traction — people committing an email", "PostHog (waitlist_signup)"],
        ["Cumulative signups vs 2,000", "Are we going to make it in time?", "Dashboard: Path to 2,000"],
        ["Engagement rate", "Is the content resonating?", "Each platform + Post Log"],
        ["Click-through rate (CTR)", "Is the post making people click the link?", "Post Log / PostHog"],
        ["Qualified visitors / week", "Top-of-funnel volume from social", "PostHog (by utm_source)"],
        ["Visitor → Signup %", "Is the website converting traffic?", "PostHog"],
    ],
    widths=[2.6, 2.7, 2.2],
)
para("Engagement rate = (likes + comments + shares + saves) ÷ impressions. "
     "CTR = link clicks ÷ impressions.", italic=True, color=GREY, size=9.5)

h2("The pacing math (the “are we on track?” number)")
para("The dashboard does this automatically, but here is the logic the founder should know:")
bullet(" 2,000 − signups so far.", bold_lead="Remaining = ")
bullet(" Remaining ÷ weeks left before launch.", bold_lead="Required per week = ")
bullet(" the average signups/week we are actually getting.", bold_lead="Run-rate = ")
bullet(" if run-rate ≥ required, we are ON PACE (green). If it is short, we are BEHIND "
       "(red) and need more volume or better conversion — now, not later.", bold_lead="Status = ")
para("Example: if launch is 8 weeks away and we have 8 signups, we need roughly 250 signups "
     "per week. The dashboard flags red the moment our weekly actual falls below that line.",
     italic=True, color=GREY, size=9.5)

h2("Quality guardrails (so we don't chase junk traffic)")
para("A channel that brings big numbers but lots of invalid emails or instant bounces is not "
     "real traction. We keep an eye on invalid-email rate and bot/rate-limit hits (both already "
     "tracked in PostHog) and discount channels that spike them.")

# ══ 3. WEEKLY RITUAL ═══════════════════════════════════════════════════════════
h1("3.  The 20-minute Monday review")
para("Same four reads, every week, in this order:")
numbered(" open ‘Path to 2,000’. Are we ON PACE or BEHIND? This is the one cell that matters most.",
         bold_lead="Pace check — ")
numbered(" signups this week vs last week on ‘Weekly Inputs’. Is the line going up?",
         bold_lead="Momentum — ")
numbered(" on ‘Channel Scorecard’, which platform produced the most signups per 1,000 "
         "impressions?", bold_lead="Best channel — ")
numbered(" pour more time into the top channel; cut or rework the worst. Decide one change "
         "for the week.", bold_lead="Act — ")

# ══ 4. PER-PLATFORM ROADMAP ════════════════════════════════════════════════════
h1("4.  Per-platform roadmap")
para("Seven channels, each with its own rhythm and the engagement signal that matters there. "
     "Always post the matching UTM link (from the dashboard’s ‘UTM Links’ tab) so PostHog can "
     "attribute the traffic.")
table(
    ["Platform", "Cadence", "Best content", "Intent signal", "Key rules"],
    [
        ["Reddit", "2–3 value posts/wk per subreddit", "Genuine help; ‘I built this’ in maker subs", "Upvotes, saves, comments", "Lead with value, link in body/comment per sub rules. Highest intent, fastest to get banned for spam."],
        ["X / Twitter", "1–2 posts/day + replies", "Build-in-public threads, before/after, polls", "Bookmarks, reposts", "Put the link in a reply or end of thread — not the first tweet."],
        ["Threads", "1–2/day", "Conversational fashion takes, casual demos", "Replies, likes", "Reach favours text + conversation; drop the link in a follow-up."],
        ["Facebook", "3–4/wk + relevant Groups", "Lifestyle/outfit posts, Group value posts", "Shares, comments", "Groups beat the Page for cold reach."],
        ["LinkedIn", "2–3/wk", "Founder journey, problem → solution", "Comments, reposts", "Already our best converter — keep it personal."],
        ["Instagram", "4–5/wk reels + stories", "Try-on, wardrobe, Aira demo reels", "Saves, shares", "Use the bio link; saves are the real intent signal. Shows as ‘ig’ in PostHog."],
        ["Product Hunt", "Launch-day spike", "Launch post + maker comments", "Upvotes", "One big moment, not weekly — line it up for the back half."],
    ],
    widths=[1.0, 1.35, 1.6, 1.0, 1.6],
)

h2("4-week launch roadmap")
table(
    ["Week", "Focus", "Goal"],
    [
        ["Week 1", "Warm up all profiles; seed pure-value posts (no hard selling). Start the Post Log.", "Establish baselines; learn what lands"],
        ["Week 2", "Ramp link posts; A/B two content styles per channel. First Monday review.", "First real signup numbers; find the top 2 channels"],
        ["Week 3", "Double down on the top 2 channels. Run the Product Hunt launch.", "A volume spike; push cumulative toward pace"],
        ["Week 4", "Scale winners, retire losers, present the trend to the founder.", "Lock the channel mix that gets us to 2,000"],
    ],
    widths=[0.8, 4.0, 2.6],
)

# ══ 5. USING THE DASHBOARD ═════════════════════════════════════════════════════
h1("5.  How to use the dashboard")
para("Companion file: Attira_Growth_Dashboard.xlsx. The rule is simple — type into the YELLOW "
     "cells, read the BLUE cells (they are formulas; don’t overwrite them).")
table(
    ["Sheet", "What it’s for", "Who touches it"],
    [
        ["README", "Instructions + colour legend", "Read once"],
        ["Post Log", "One row per post: likes, comments, shares, saves, impressions", "Fill after every post"],
        ["Weekly Inputs", "Weekly visitors + signups from PostHog", "Fill once a week"],
        ["Funnel Dashboard", "All-time funnel + conversion rates", "Read"],
        ["Path to 2,000", "Cumulative vs goal, pace, run-rate, status", "Read — the headline"],
        ["Channel Scorecard", "Per-platform comparison", "Paste signups-by-source weekly; read the rest"],
        ["UTM Links", "The links to actually share", "Copy links from here"],
        ["Targets", "Goal, launch date, KPI bars", "Set once, adjust as needed"],
    ],
    widths=[1.4, 3.6, 2.0],
)
para("Three inputs keep it alive: (1) a Post Log row after each post, (2) weekly visitors + "
     "signups from PostHog, (3) weekly signups-by-source pasted into the Channel Scorecard. "
     "That’s it — everything else calculates itself.", italic=True, color=GREY, size=9.5)

# ══ 6. PHASE 2 ═════════════════════════════════════════════════════════════════
h1("6.  Phase 2 — beta tracking (once we share the app by link)")
para("In about two weeks we start distributing the iOS beta via a public Apple TestFlight "
     "link. At that point we extend the funnel and switch the North Star to Activated beta "
     "users / week. Planned changes (already scoped, not built yet):")
bullet(" a tracked ‘Get beta access’ button appears after a visitor joins the waitlist.",
       bold_lead="Website: ")
bullet(" a beta_access_clicked event in PostHog captures who clicked and from which source, "
       "before they leave for TestFlight.", bold_lead="Analytics: ")
bullet(" Apple hides install attribution, so TestFlight installs and active testers get "
       "entered weekly by hand from App Store Connect. We will state this limitation plainly — "
       "the click event is our best on-site signal.", bold_lead="Honest caveat: ")
para("The dashboard already has hidden columns reserved for these numbers, so turning Phase 2 "
     "on is a small step, not a rebuild.", italic=True, color=GREY, size=9.5)

# ══ 7. TARGETS ═════════════════════════════════════════════════════════════════
h1("7.  What ‘good’ looks like")
table(
    ["Checkpoint", "Signal we want to see"],
    [
        ["Week 1", "Baselines set; engagement rate ≥ 3–5% on at least one channel"],
        ["Week 2", "First signups flowing; Visitor → Signup ≥ 15%; a clear top-2 channel"],
        ["Week 4", "Weekly signups at or above the required pace; cumulative trending to 2,000"],
        ["Pre-launch", "≥ 2,000 cumulative waitlist signups"],
    ],
    widths=[1.6, 5.4],
)
para("The North Star is not a single week’s number — it’s staying on or above the pace line, "
     "week after week, until we cross 2,000.", italic=True, color=GREY)

out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "Attira_Growth_Strategy.docx")
doc.save(out)
print("Wrote", out)
